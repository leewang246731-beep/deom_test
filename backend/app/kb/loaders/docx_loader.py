"""DOCX Loader — using python-docx with heading preservation."""
from app.kb.loaders.base_loader import BaseDocumentLoader, LoadedDocument


class DocxLoader(BaseDocumentLoader):
    @staticmethod
    def supported_extensions() -> list[str]:
        return [".docx"]

    def load(self, file_path: str) -> LoadedDocument:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = Document(file_path)
        paragraphs = []
        current_heading = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Preserve heading hierarchy
            if para.style.name.startswith("Heading"):
                current_heading = text
                paragraphs.append(f"[{text}]")
            elif current_heading:
                paragraphs.append(f"[{current_heading}] {text}")
            else:
                paragraphs.append(text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        full_text = "\n".join(paragraphs)

        return LoadedDocument(
            text=full_text,
            metadata={
                "source": file_path,
                "format": "docx",
                "paragraph_count": len(doc.paragraphs),
            },
        )
